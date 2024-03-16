from docx import Document

#명세서 파일 이름 ex : 'P24091_청구항 5개_원본'
spec_file_name = input('명세서 파일의 이름을 확장자 *없이* 입력하세요 : ')
expand = '.docx'

#1. 명세서 복사본 생성해서 도면 제거하기
spec_orig = Document(spec_file_name+expand)
spec_new = Document()
text_arr = []

#대표도, 도면, 도 등 삭제, 도 7 넘으면 추가할 것
for line in spec_orig.paragraphs:
    if '【대표도】' not in line.text:
        if '【도면】' not in line.text:
            if '【도 1】' not in line.text:
                if '【도 2】' not in line.text:
                    if '【도 3】' not in line.text:
                        if '【도 4】' not in line.text:
                            if '【도 5】' not in line.text:
                                if '【도 6】' not in line.text:
                                    if '【도 7】' not in line.text:
                                        text_arr.append(line.text)

for i in range(len(text_arr)):
    spec_new.add_paragraph(text_arr[i])

spec_new.save(spec_file_name+'_GPT업로드용'+expand)



